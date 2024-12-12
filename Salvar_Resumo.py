from fpdf import FPDF
from docx import Document

class SaveSummaries:
    """salva resumos em PDF e Word e txt"""
    
    @staticmethod
    def save_to_pdf(summaries, output_file="summary.pdf"):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        pdf.cell(200, 10, txt="Resumo do Vídeo", ln=True, align='C')
        pdf.ln(10)  # Adiciona uma linha em branco

        for summary in enumerate(summaries, start=1):
            pdf.multi_cell(0, 10, txt=f"- {summary}", align='L')
            pdf.ln(5)  # Espaçamento entre os resumos

        pdf.output(output_file)
        print(f"Resumo salvo em PDF: {output_file}")

    @staticmethod
    def save_to_word(summaries, output_file="summary.docx"):
        """Salva os resumos em um arquivo Word."""
        document = Document()
        document.add_heading("Resumo do Vídeo", level=1)

        for summary in enumerate(summaries, start=1):
            document.add_paragraph(f"- {summary}")

        document.save(output_file)
        print(f"Resumo salvo em Word: {output_file}")

    @staticmethod
    def save_to_txt(summaries, output_file="summary.txt"):
        """Salva os resumos em um arquivo TXT."""
        with open(output_file, 'w') as f:
            f.write("Resumo do Vídeo\n\n")
            for summary in enumerate(summaries, start=1):
                f.write(f"- {summary[1]}\n")
        print(f"Resumo salvo em TXT: {output_file}")
